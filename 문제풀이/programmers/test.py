# -*- coding: utf-8 -*-
"""
Created on Fri Aug 29 14:08:11 2025

@author: mindo
"""

from docx import Document

# 새 문서 생성
doc = Document()

# 상단 좌측 표
table_top_left = doc.add_table(rows=4, cols=2)
table_top_left.style = "Table Grid"
table_top_left.cell(0,0).text = "문서번호"
table_top_left.cell(0,1).text = ""
table_top_left.cell(1,0).text = "시행일자"
table_top_left.cell(1,1).text = "'20.09.01"
table_top_left.cell(2,0).text = "문서구분"
table_top_left.cell(2,1).text = "비공개"
table_top_left.cell(3,0).text = "열람등급"
table_top_left.cell(3,1).text = "Ⅹ"

# 제목
doc.add_paragraph("")
doc.add_heading("차 행성 침공계획 [보고]", level=1)

# 상단 우측 결재란 표
table_top_right = doc.add_table(rows=5, cols=5)
table_top_right.style = "Table Grid"
top_titles = ["기안", "검토", "심사", "결재", "결재"]
for i, title in enumerate(top_titles):
    table_top_right.cell(0,i).text = title

roles = ["작성자", "실무자", "부서장", "법리검토", "법리판단"]
for i, role in enumerate(roles):
    table_top_right.cell(1,i).text = role
    for r in range(2,5):
        table_top_right.cell(r,i).text = ""

# 본문
doc.add_heading("□ 개요 및 목적", level=2)
doc.add_paragraph("○ 차 행성 침공으로 저그 위협 봉쇄")
doc.add_paragraph("○ 기지 구축으로 자치령 영역 확대")
doc.add_paragraph("○ 중추석을 이용한 칼날 여왕 제거")
doc.add_paragraph(" * 참고 교전장: 종교전 녹색 11")

doc.add_heading("□ 일시", level=2)
doc.add_paragraph("○ 작전 일시 : 2504. 10. 01 ~")
doc.add_paragraph("○ 작전 구역 : 차 행성, 차 알레프 제도 정거장")

# 세부계획 표
doc.add_heading("□ 세부계획", level=2)
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "구분"
hdr[1].text = "자치령 무공집단"
hdr[2].text = "자치령 해병단"
hdr[3].text = "기타"

row = table.add_row().cells
row[0].text = "동원부대"
row[1].text = "제1함대"
row[2].text = "제1해병사단 (병력 14577명)"
row[3].text = "베헤모스급 전투순양함 564척"

# 동원 병력
doc.add_heading("□ 동원 병력 (세부)", level=2)
doc.add_paragraph("○ 제1함대 (병력 4753명)")
doc.add_paragraph(" - 기갑 부대 : 부세팔루스 외 미노타우르스크 전투순양함 25척")
doc.add_paragraph(" - 군수지원함")
doc.add_paragraph(" - 1함대 항공단")
doc.add_paragraph(" - 메딕버스 85발")
doc.add_paragraph("○ 제1해병사단 (병력 14577명)")
doc.add_paragraph(" - 1해병여단")
doc.add_paragraph(" - 2해병여단")
doc.add_paragraph(" - 3해병여단")
doc.add_paragraph(" - 4해병여단")
doc.add_paragraph(" - 5해병여단")
doc.add_paragraph(" - 6해병여단")
doc.add_paragraph(" - 19기갑여단")
doc.add_paragraph("○ 제7특공대대 (병력 564명)")
doc.add_paragraph("○ 베헤모스급 전투순양함 헤피리온 1척")

# 행정사항
doc.add_heading("□ 행정사항", level=2)
doc.add_paragraph("○ 중추석 직공선으로 보호")
doc.add_paragraph("○ 차 알레프 공역 내 한때 밀집배열 전개")
doc.add_paragraph("○ 행방불명 차원봉 4133명 확인 중")
doc.add_paragraph("○ 한때는 자폭선 방패 생명체제 우익할 것")
doc.add_paragraph("○ 해병사단 병무교본 및 무기 공인, 군사법 무효 시 병공망 제한")
doc.add_paragraph("○ 33해병여단 예비대 특공대 투입 후 37해병여단 합류")
doc.add_paragraph("○ 병력 이동시 대공포 보강 및 중추석 교환 필수")
doc.add_paragraph("○ 각 부대 인원에게 제로투 투입계획 적용, 중수석 가동")
doc.add_paragraph("○ 중추석 예비 교환인원 배치")
doc.add_paragraph("○ 전투종결 및 칼날여왕 제거 후 기지 구축")

# 파일 저장
doc.save("차_행성_침공계획_양식.docx")
print("완료: 차_행성_침공계획_양식.docx")